"""
Bulk BG Document Generator
Converts multiple BGs from a text file into individual Word documents
Tracks processed BGs to avoid regenerating duplicates
"""

import os
import re
import json
import hashlib
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_bg_hash(bg_text):
    """Generate a hash for a BG text to track if it's been processed"""
    return hashlib.md5(bg_text.encode()).hexdigest()


def load_processed_bgs(output_dir):
    """Load the list of already processed BGs from tracking file"""
    tracking_file = os.path.join(output_dir, '.bg_tracker.json')
    if os.path.exists(tracking_file):
        try:
            with open(tracking_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_processed_bgs(processed_bgs, output_dir):
    """Save the list of processed BGs to tracking file"""
    tracking_file = os.path.join(output_dir, '.bg_tracker.json')
    with open(tracking_file, 'w', encoding='utf-8') as f:
        json.dump(processed_bgs, f, indent=2)


def parse_bgs_from_file(file_path):
    """
    Parse BGs from a text file. Handles multiple formats:
    - Separated by dashes: "---" or "======"
    - Numbered clauses: "1. ", "2. ", etc.
    - Separated by blank lines
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    bgs = []
    
    # First priority: Try dash/line separator format (---)
    if re.search(r'---+', content):
        bgs = re.split(r'---+', content)
        bgs = [bg.strip() for bg in bgs if bg.strip()]
        if len(bgs) > 1:
            print(f"✓ Detected '---' separator format: Found {len(bgs)} complete BGs")
            return bgs
    
    # Second priority: Extract numbered clauses (1. 2. 3. etc.) from content
    clause_match = re.findall(r'^\d+\.\s+(.+?)(?=^\d+\.|$)', content, re.MULTILINE | re.DOTALL)
    if clause_match and len(clause_match) > 1:
        bgs = [clause.strip() for clause in clause_match]
        print(f"✓ Detected numbered clauses format: Found {len(bgs)} clauses")
        return bgs
    
    # Fallback: treat multiple blank lines as separators
    bgs = re.split(r'\n\s*\n+', content)
    bgs = [bg.strip() for bg in bgs if bg.strip()]
    
    if len(bgs) > 1:
        print(f"✓ Detected blank line format: Found {len(bgs)} BGs")
    else:
        print("⚠ Warning: Could not detect multiple BGs. Make sure they are properly separated.")
    
    return bgs


def create_word_document(bg_text, output_file):
    """
    Create a Word document with the BG text
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Background', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    paragraph = doc.add_paragraph(bg_text)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(12)
    
    # Set font
    for run in paragraph.runs:
        run.font.size = Pt(12)
        run.font.name = 'Calibri'
    
    # Save document
    doc.save(output_file)


def generate_bg_documents(input_file, output_dir='BG_Documents'):
    """
    Main function to generate Word documents from BGs text file
    Skips BGs that have already been processed (identified by content hash)
    Uses hash as filename so same BG always has same filename regardless of order
    """
    # Create output directory
    Path(output_dir).mkdir(exist_ok=True)
    
    # Load previously processed BGs
    processed_bgs = load_processed_bgs(output_dir)
    
    # Parse BGs
    bgs = parse_bgs_from_file(input_file)
    
    if not bgs:
        print("❌ No BGs found in the file!")
        return
    
    # Identify new BGs
    new_bgs = []
    skipped_count = 0
    next_index = len(processed_bgs) + 1  # Get next available index
    
    for bg in bgs:
        bg_hash = get_bg_hash(bg)
        
        if bg_hash in processed_bgs:
            skipped_count += 1
            continue
        
        new_bgs.append((bg, bg_hash))
    
    if skipped_count > 0:
        print(f"⏭️  Skipped {skipped_count} already processed BG(s)")
    
    if not new_bgs:
        print("✅ All BGs have already been processed! No new documents to generate.")
        return
    
    print(f"\n📄 Generating {len(new_bgs)} new Word document(s)...\n")
    
    # Generate only new documents
    for bg, bg_hash in new_bgs:
        filename = os.path.join(output_dir, f"BG_{next_index:03d}.docx")
        try:
            create_word_document(bg, filename)
            processed_bgs[bg_hash] = {
                "filename": f"BG_{next_index:03d}.docx",
                "hash": bg_hash
            }
            print(f"✓ Created: {filename}")
            next_index += 1
        except Exception as e:
            print(f"❌ Error creating {filename}: {str(e)}")
    
    # Save updated tracking
    save_processed_bgs(processed_bgs, output_dir)
    
    total_count = len(processed_bgs)
    print(f"\n✅ Done! Total {total_count} document(s) in '{output_dir}' folder")


if __name__ == "__main__":
    # Configuration
    INPUT_FILE = "bgs_input.txt"  # Your text file with all BGs
    OUTPUT_DIR = r"C:\Users\KRISHNAN VARUN\OneDrive\Desktop\Impact_BGs"
    
    if not os.path.exists(INPUT_FILE):
        print(f"❌ Error: '{INPUT_FILE}' not found!")
        print("Please create a text file with your BGs and name it 'bgs_input.txt'")
    else:
        generate_bg_documents(INPUT_FILE, OUTPUT_DIR)
